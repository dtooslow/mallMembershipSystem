import docx
from docx.shared import Inches
import os

def build_document():
    doc = docx.Document()
    
    doc.add_heading('Mall Membership System - Project Documentation', 0)
    
    # --- VI. ERD ---
    doc.add_heading('VI. Entity Relationship Diagram (ERD)', 1)
    doc.add_paragraph("The database structure of the Mall Membership System handles customers, administrative permissions, merchant shops, product listings, loyalty points-based transactions, reward definitions, reward redemptions, and system-level notifications.")
    doc.add_heading('1. Database Entities & Attributes', 2)
    doc.add_paragraph("Users, Admins, Memberships, Shops, Products, Transactions, Rewards, Reward Redemptions, Notifications, Events.")
    doc.add_heading('2. Relationships Model Diagram', 2)
    
    erd_path = os.path.join("scratch", "screenshots", "erd_diagram.png")
    if os.path.exists(erd_path):
        doc.add_picture(erd_path, width=Inches(6.0))
    else:
        doc.add_paragraph("[ERD Image Missing]")
    
    # --- VII. User Interface Manual ---
    doc.add_heading('VII. User Interface Manual', 1)
    doc.add_paragraph("This manual provides a detailed step-by-step user guide for customer members and system administrators.")
    doc.add_heading('1. Key Navigation Flows', 2)
    doc.add_paragraph("Guest Navigation & Landing Page: Access the root URL. Browse shops and rewards.")
    doc.add_paragraph("Customer Portal Flow: Register, Login, Apply for Membership, View Dashboard, Purchase Products, Claim Rewards.")
    doc.add_paragraph("Administrator Portal Flow: Admin Login, Dashboard Statistics, Moderate Memberships, Manage Shops, Manage Rewards.")
    
    # --- VIII. Gantt Chart ---
    doc.add_heading('VIII. Gantt Chart', 1)
    doc.add_paragraph("The project development was executed over a 4-week lifecycle involving requirements, database setups, UI styling, MVC implementation, and integration testing.")
    gantt_path = os.path.join("scratch", "screenshots", "gantt_chart.png")
    if os.path.exists(gantt_path):
        doc.add_picture(gantt_path, width=Inches(6.0))
    else:
        doc.add_paragraph("[Gantt Chart Image Missing]")
    
    # --- IX. UI/UX ---
    doc.add_heading('IX. UI/UX', 1)
    doc.add_paragraph("Below are high-quality visual representations of each main application interface.")
    
    screenshots = [
        ("1. Public Landing Page", "welcome.png"),
        ("2. Admin Login Page", "admin_login.png"),
        ("3. Admin Dashboard", "admin_dashboard.png"),
        ("4. Admin Member Directory", "admin_memberships.png"),
        ("5. Admin Manage Shops", "admin_shops.png"),
        ("6. Admin Manage Rewards", "admin_rewards.png"),
        ("7. Admin Transactions", "admin_transactions.png"),
        ("8. Admin Events", "admin_events.png"),
        ("9. Customer Login", "user_login.png"),
        ("10. Customer Registration", "user_register.png"),
        ("11. Member Dashboard", "user_dashboard.png"),
        ("12. Apply for Membership", "membership_apply.png")
    ]
    
    for title, filename in screenshots:
        doc.add_heading(title, 2)
        img_path = os.path.join("scratch", "screenshots", filename)
        if os.path.exists(img_path):
            doc.add_picture(img_path, width=Inches(6.0))

    output_path = "Project_Documentation_Final.docx"
    doc.save(output_path)
    print(f"Document successfully saved to {output_path}")

if __name__ == "__main__":
    build_document()
