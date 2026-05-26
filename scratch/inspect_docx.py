import docx
import sys

sys.stdout.reconfigure(encoding='utf-8')

doc = docx.Document("project_documentation_sections_VI_IX.docx")

def print_range(start, end):
    print(f"--- Paragraphs {start} to {end} ---")
    for idx in range(start, min(end, len(doc.paragraphs))):
        p = doc.paragraphs[idx]
        print(f"P{idx} (style: {p.style.name}): {repr(p.text)}")

print_range(30, 60)
print_range(160, 220)
